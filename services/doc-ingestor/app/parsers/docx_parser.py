"""
DOCX Parser - Extract text and metadata from Word documents
"""

from docx import Document as DocxDocument
from typing import Dict, Any, Tuple
import structlog

logger = structlog.get_logger()


class DOCXParser:
    """DOCX document parser"""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse DOCX file and extract text and metadata
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            # Extract metadata
            metadata = {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
                "last_modified_by": doc.core_properties.last_modified_by,
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error("Error parsing DOCX", file=file_path, error=str(e))
            raise
